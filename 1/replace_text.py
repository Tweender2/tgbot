from docx import Document
from datetime import datetime
import locale
import re
from num2words import num2words
import os


locale.setlocale(locale.LC_TIME, 'Russian')

def replace_text_in_docx(doc_path, old_info_text, new_info_text, pattern):
    doc = Document(f"{doc_path}.docx")
    
    # Создаем словарь для замены
    replacements = dict(zip(old_info_text, new_info_text))
    print(replacements)

    for key, value in replacements.items():
        if key == 'Price':
            # Обработка цены
            try:
                # Удаляем все символы, кроме цифр и запятой/точки
                cleaned_value = re.sub(r'[^\d.,]', '', value)
                price_value = float(cleaned_value.replace(',', '.'))  # Обработка ввода с запятой
                rubles = int(price_value)
                kopecks = int((price_value - rubles) * 100)
                formatted_price = f"{rubles:,} руб. ({num2words(rubles, lang='ru')} рублей {kopecks:02d} копеек)".replace(',', ' ')

                for paragraph in doc.paragraphs:
                    for run in paragraph.runs:
                        if key in run.text:
                            print(f"Заменяем {key} на '{formatted_price}' в параграфе")
                            run.text = run.text.replace(key, formatted_price)

                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    if key in run.text:
                                        print(f"Заменяем {key} на '{formatted_price}' в таблице")
                                        run.text = run.text.replace(key, formatted_price)
            except ValueError:
                print("Ошибка: введено некорректное значение для цены.")

    # Для заполнения ОГРНИП в тексте
    for key, value in replacements.items():
        if key == '<ОГРНИП>':
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    if "<ОГРНИПHA>" in run.text:
                        print(f"Заменяем '<ОГРНИПHA>' на '{value}' в параграфе")
                        run.text = run.text.replace("<ОГРНИПHA>", value)

    # Для заполнения названия организации
    if len(old_info_text) > len(new_info_text):
        replacements['description'] = 'Пустота'
    for key, value in replacements.items():
        if key == 'description' and value == 'Пустота' or key == 'description' and value == '':
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                if key in run.text:
                                    print(f"Заменяем {key} на '' в таблице")
                                    run.text = run.text.replace(key, '')
                                    if paragraph.text.strip() == '':
                                        p = paragraph._element
                                        p.getparent().remove(p)
                                        print("Выполнено удаление :)")
                                        break


    for paragraph in doc.paragraphs:
        for old_text, new_text in replacements.items():
            for run in paragraph.runs:
                if old_text in run.text:
                    if pattern == 'Choice1' and old_text == 'description' and new_text != '':
                        print(f"Заменяем {old_text} на ' ({new_text})' в параграфе")
                        run.text = run.text.replace(old_text, f' ({new_text})')
                    else:
                        print(f"Заменяем {old_text} на '{new_text}' в параграфе")
                        run.text = run.text.replace(old_text, new_text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for old_text, new_text in replacements.items():
                        for run in paragraph.runs:
                            if old_text in run.text:
                                if pattern == 'Choice1' and old_text == 'description' and new_text != '':
                                        print(f"Заменяем {old_text} на '({new_text})' в таблице")
                                        run.text = run.text.replace(old_text, f'({new_text})')
                                        if paragraph.text.strip() == '':
                                            p = paragraph._element
                                            p.getparent().remove(p)
                                            print("Выполнено удаление :)")
                                            break
                                else:
                                    print(f"Заменяем {old_text} на '{new_text}' в таблице")
                                    run.text = run.text.replace(old_text, new_text)

    
                            
    # Замена особым способом для инициалов
    for key, value in replacements.items():
        if key == 'full_name' or key == '<ФИО>' or key == 'FIO':
            initials = 'name'
            parts = value.split()
            last_name = parts[0]
            first_name = parts[1][0] + '.'
            patronymic = parts[2][0] + '.'
            fio = '%s %s%s' % (last_name, first_name, patronymic)
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    if initials in run.text:
                        print(f"Заменяем {initials} на '{fio}' в параграфе")
                        run.text = run.text.replace(initials, fio)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                if initials in run.text:
                                    print(f"Заменяем {initials} на '{fio}' в таблице")
                                    run.text = run.text.replace(initials, fio)


    
                                    

    # Замена особым способом для дат
    for key, value in replacements.items():
        if key == 'full_date' or key == '<Дата>' or key == 'date':
            old_date = 'date'
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    if old_date in run.text:
                        print(f"Заменяем {old_date} на '{value}' в параграфе")
                        run.text = run.text.replace(old_date, value)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                if old_date in run.text:
                                    print(f"Заменяем {old_date} на '{value}' в таблице")
                                    run.text = run.text.replace(old_date, value)

    # Замена особым способом для второй даты
    for key, value in replacements.items():
        if key == 'data' or key == '<Атад>':
            old_date = 'data'
            try:
                for paragraph in doc.paragraphs:
                    for run in paragraph.runs:
                        if old_date in run.text:
                            print(f"Заменяем {old_date} на '{value}' в параграфе")
                            run.text = run.text.replace(old_date, value)
                
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    if old_date in run.text:
                                        print(f"Заменяем {old_date} на '{value}' в таблице")
                                        run.text = run.text.replace(old_date, value)
            except ValueError:
                print("Ошибка: неверный формат даты. Ожидается 'дд.мм.гггг'.")
    
    # Замена особым способом для третей даты
    for key, value in replacements.items():
        if key == 'entry' or key == '<Энтри>':
            old_date = 'entry'
            try:
                for paragraph in doc.paragraphs:
                    for run in paragraph.runs:
                        if old_date in run.text:
                            print(f"Заменяем {old_date} на '{value}' в параграфе")
                            run.text = run.text.replace(old_date, value)
                
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    if old_date in run.text:
                                        print(f"Заменяем {old_date} на '{value}' в таблице")
                                        run.text = run.text.replace(old_date, value)
            except ValueError:
                print("Ошибка: неверный формат даты. Ожидается 'дд.мм.гггг'.")

    # Замена на сокращения договора
    for key, value in replacements.items():
        new_post = ''
        score = 1
        index = 0
        post = 'position'
        separation = value.split()
        while score <= len(separation):
            letter = separation[index][0]
            new_post += letter.upper()
            score += 1
            index += 1
        if key == '<Должность>':
            if len(value.split()) > 1:
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    if post in run.text:
                                        print(f"Заменяем {post} на '{new_post}' в таблице")
                                        run.text = run.text.replace(post, new_post)
            elif len(value.split()) == 1:
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    if post in run.text:
                                        print(f"Заменяем {post} на '{value}' в таблице")
                                        run.text = run.text.replace(post, value)
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    if post in run.text:
                        print(f"Заменяем {post} на '{value}' в параграфе")
                        run.text = run.text.replace(post, value)

    # Замена на в лице должность
    for key, value in replacements.items():
        if key == 'post' or key == 'OST':
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    if key in run.text:
                        print(f"Заменяем {key} на '{value}' в параграфе")
                        run.text = run.text.replace(key, value)
    
    # Сохраняем измененный документ
    num = ''
    if "NumberAgreement" in replacements:
        num = replacements.get("NumberAgreement")
    elif "<Номер договора>" in replacements:
        num = replacements.get("<Номер договора>")

    arg = ''
    if "Agr" in replacements:
        arg = replacements.get("Agr")
    
    fio_name = ''
    if "FIO" in replacements:
        fio_name = replacements.get("FIO")
        familia = fio_name.split(' ')

    if str(pattern) in ['Choice_1']:
        print(f"\nИП {familia[0]} доп. соглашение {num} к договору {arg}.docx\n")
        doc.save(f"ИП {familia[0]} доп. соглашение {num} к договору {arg}.docx")
        os.remove(f"{doc_path}.docx")
        return f"ИП {familia[0]} доп. соглашение {num} к договору {arg}.docx"
    elif str(pattern) in ['Choice1']:
        for key, value in replacements.items():
            if key == "<ФИО>":
                name = value
        print(f"\nPattern/ИП {name} договор размещения рекламы {num}.docx\n")
        doc.save(f"Pattern/ИП {name} договор размещения рекламы {num}.docx")
        os.remove(f"{doc_path}.docx")
        return f"Pattern/ИП {name} договор размещения рекламы {num}.docx"
    else:
        organization = ''
        if "description" in replacements:
            organization = replacements.get("description")
        elif "Organization" in replacements:
            organization = replacements.get("Organization")   
    
        print(organization)
        if (organization in ['', 'Пустота'] and str(pattern) in ['Choice2']) or (
            organization in ['', 'Пустота'] and str(pattern) in ['Choice_2']):
            # Случай 1: ничего не вписано в 17 строчку
            doc_name = str(doc_path).split("/")
            print(f"\nPattern/{doc_name[1]} {num}.docx\n")
            doc.save(f"Pattern/{doc_name[1]} {num}.docx")
            os.remove(f"{doc_path}.docx")
            return f"Pattern/{doc_name[1]} {num}.docx"
        elif ("«" in str(organization) and str(pattern) in ['Choice2']) or (
            "«" in str(organization) and str(pattern) in ['Choice_2']):
            # Случай 3: текст со скобочками
            org = str(organization).split("«")
            organiz = str(org[1]).split("»")
            org_name = organiz[0].strip()
            
            doc_name = str(doc_path).split("/")
            print(f"\nPattern/ООО_{org_name}_{doc_name[1]} {num}.docx\n")
            doc.save(f"Pattern/ООО_{org_name}_{doc_name[1]} {num}.docx")
            os.remove(f"{doc_path}.docx")
            return f"Pattern/ООО_{org_name}_{doc_name[1]} {num}.docx"
        elif str(pattern) in ['Choice2']:
            # Случай 2: простой текст без скобок
            org_name = str(organization).strip()
            
            doc_name = str(doc_path).split("/")
            print(f"\nPattern/ООО_{org_name}_{doc_name[1]} {num}.docx\n")
            doc.save(f"Pattern/ООО_{org_name}_{doc_name[1]} {num}.docx")
            os.remove(f"{doc_path}.docx")
            return f"Pattern/ООО_{org_name}_{doc_name[1]} {num}.docx"
        elif str(pattern) in ['Choice_2']:
            org_name = str(organization).strip()
            
            print(f"\nPattern/ООО_{org_name}_доп. соглашение {num} к договору {arg}.docx\n")
            doc.save(f"Pattern/ООО_{org_name}_доп. соглашение {num} к договору {arg}.docx")
            os.remove(f"{doc_path}.docx")
            return f"Pattern/ООО_{org_name}_доп. соглашение {num} к договору {arg}.docx"
        
        if str(pattern).startswith("Pattern"):
            print(f"\n{doc_path} {num}.docx\n")
            doc.save(f"{doc_path} {num}.docx")
            os.remove(f"{doc_path}.docx")
            return f"{doc_path} {num}.docx"